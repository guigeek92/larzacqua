import copy
import os
import re
import sys
import zipfile
from typing import Dict, Iterable, List, Tuple

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches
from docx2pdf import convert
# Ajouter le dossier parent pour accéder aux modules internes
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

# Import des modules du projet
from modules.loader import load_data
from modules.hydraulics import compute_hydraulics
from modules.power import compute_power
from modules.scoring import score_sites
from modules.turbine import load_turbine_db
from modules.turbine_selector import select_turbine
from modules import productible


# -------------------- CONFIG -------------------- #

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DEFAULT_TEMPLATE = os.path.join(BASE_DIR, '..', 'template.docx')
PROVIDED_TEMPLATE = os.path.join(BASE_DIR, '..', 'outputs', 'rapport_temp.docx')
OUTPUT_DIR = os.path.join(BASE_DIR, '..', 'outputs')
TEMP_DOCX = os.path.join(OUTPUT_DIR, 'rapport_temp_gen.docx')
OUTPUT_PDF = os.path.join(OUTPUT_DIR, 'rapport_final.pdf')
CSV_PATH = os.path.join(BASE_DIR, '..', 'CSV', 'aep_organe_pression.csv')
TURBINE_DB_PATH = os.path.join(BASE_DIR, '..', 'CSV', 'turbine_db.csv')
TOP_N_SITES = 5


# -------------------- UTILS -------------------- #

def replace_placeholders_in_paragraph(paragraph, variables: Dict[str, str]) -> None:
    """Remplace les placeholders sans écraser les images ni le style."""
    for run in paragraph.runs:
        text = run.text
        for key, val in variables.items():
            text = text.replace(f'{{{{ {key} }}}}', str(val)) \
                       .replace(f'{{{{{key}}}}}', str(val)) \
                       .replace(f'{{ {key} }}', str(val)) \
                       .replace(f'{{{key}}}', str(val))
        if text != run.text:
            run.text = text


def strip_toc_errors(paragraph) -> None:
    """Supprime le texte d'erreur du sommaire s'il apparait."""
    if 'Erreur ! Signet non defini.' in paragraph.text:
        for run in paragraph.runs:
            run.text = run.text.replace('Erreur ! Signet non defini.', '')


def set_table_column_widths(table: Document, widths_in: List[float]) -> None:
    """Applique des largeurs fixes aux colonnes d'un tableau."""
    if not widths_in:
        return
    for col_idx, width in enumerate(widths_in):
        for row in table.rows:
            if col_idx < len(row.cells):
                row.cells[col_idx].width = Inches(width)


def set_cell_text(cell, value) -> None:
    """Écrit la valeur dans la cellule en conservant le style des paragraphes."""
    text = str(value)
    if not cell.paragraphs:
        cell.text = text
        return
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''
    cell.paragraphs[0].add_run(text)


def has_placeholder(text: str, key: str) -> bool:
    return f"{{{{{key}}}}}" in text.replace(' ', '')


def clear_paragraph_text(paragraph) -> None:
    if paragraph.runs:
        for run in paragraph.runs:
            run.text = ''
    else:
        paragraph.text = ''


def docx_has_images(docx_path: str) -> bool:
    """Détecte si le docx référence des images (drawing, pict, VML)."""
    if not os.path.exists(docx_path):
        return False
    try:
        with zipfile.ZipFile(docx_path) as z:
            for name in z.namelist():
                if not name.endswith('.xml'):
                    continue
                xml = z.read(name).decode('utf-8', errors='ignore')
                if '<w:drawing' in xml or '<w:pict' in xml or 'v:imagedata' in xml:
                    return True
    except zipfile.BadZipFile:
        return False
    return False


def strip_vml_shapes_in_docx(docx_path: str) -> None:
    """Supprime les formes VML non-image pour eviter les aplats de couleur en PDF."""
    shape_pattern = re.compile(r'<v:(?:shape|rect)\b[^>]*?(?:/>|>.*?</v:(?:shape|rect)>)', re.DOTALL)

    def remove_shape(match):
        shape = match.group(0)
        if 'imagedata' in shape or 'a:blip' in shape:
            return shape
        return ''

    temp_path = f"{docx_path}.tmp"
    with zipfile.ZipFile(docx_path, 'r') as zin, zipfile.ZipFile(temp_path, 'w') as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.startswith('word/') and item.filename.endswith('.xml'):
                if any(part in item.filename for part in ['document', 'header', 'footer']):
                    text = data.decode('utf-8', errors='ignore')
                    text = shape_pattern.sub(remove_shape, text)
                    data = text.encode('utf-8')
            zout.writestr(item, data)
    os.replace(temp_path, docx_path)


def strip_comments_in_docx(docx_path: str) -> None:
    """Supprime les commentaires et leurs ancres pour ne pas les afficher en PDF."""
    def remove_comments_xml(text):
        text = re.sub(r'<w:commentRangeStart[^>]*/>', '', text)
        text = re.sub(r'<w:commentRangeEnd[^>]*/>', '', text)
        text = re.sub(r'<w:commentReference[^>]*/>', '', text)
        return text

    temp_path = f"{docx_path}.tmp"
    with zipfile.ZipFile(docx_path, 'r') as zin, zipfile.ZipFile(temp_path, 'w') as zout:
        for item in zin.infolist():
            # Drop comments part entirely
            if item.filename == 'word/comments.xml':
                continue
            data = zin.read(item.filename)
            if item.filename.startswith('word/') and item.filename.endswith('.xml'):
                text = data.decode('utf-8', errors='ignore')
                if 'comment' in text:
                    text = remove_comments_xml(text)
                data = text.encode('utf-8')
            zout.writestr(item, data)
    os.replace(temp_path, docx_path)


def propose_turbines(site_row: pd.Series, turbine_db: pd.DataFrame, top_n: int = 3) -> pd.DataFrame:
    pressure = site_row['delta_p']
    flow = site_row['estimated_flow']
    diameter = site_row.get('diameter', 0)
    pressure_tol = 0.2 * pressure if pressure > 0 else 0.2
    flow_tol = 0.2 * flow if flow > 0 else 0.2
    diameter_tol = 20
    details = []
    for _, t in turbine_db.iterrows():
        compatible = (
            (t['pression_min_bar'] <= pressure + pressure_tol) and
            (t['pression_max_bar'] >= pressure - pressure_tol) and
            (t['debit_min_m3h'] <= flow + flow_tol) and
            (t['debit_max_m3h'] >= flow - flow_tol) and
            (t['diametre_mm'] <= diameter + diameter_tol) and
            (t['diametre_mm'] >= diameter - diameter_tol)
        )
        details.append(compatible)
    candidates = turbine_db[details].copy()
    return candidates.head(top_n)


def insert_dataframe_to_word(doc: Document, dataframe: pd.DataFrame, placeholder_key: str) -> None:
    """
    Insère un DataFrame à l'emplacement d'un placeholder dans un Word.
    Cherche dans tous les paragraphes de toutes les cellules.
    """
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Cherche le placeholder dans chaque paragraphe de la cellule
                for para in cell.paragraphs:
                    if has_placeholder(para.text, placeholder_key):
                        # Efface le placeholder
                        clear_paragraph_text(para)
                        
                        # Conserve une ligne modèle si elle existe pour préserver le style
                        template_row = table.rows[1] if len(table.rows) > 1 else None
                        while len(table.rows) > 2:
                            table._tbl.remove(table.rows[2]._tr)
                        
                        # Conserve l'en-tête existant du template
                        
                        # Ajoute les lignes du DataFrame
                        for i, (_, row_df) in enumerate(dataframe.iterrows()):
                            if i == 0 and template_row is not None:
                                row_cells = template_row.cells
                            else:
                                if template_row is not None:
                                    new_tr = copy.deepcopy(template_row._tr)
                                    table._tbl.append(new_tr)
                                    row_cells = table.rows[-1].cells
                                else:
                                    row_cells = table.add_row().cells
                            for j, value in enumerate(row_df):
                                if j < len(row_cells):
                                    set_cell_text(row_cells[j], value)
                        return

def insert_synthese_sites_table(doc: Document, results_sorted: pd.DataFrame, placeholder_key: str) -> None:
    """Insère le tableau synthèse des sites à l'emplacement d'un placeholder."""
    # Cherche dans les tableaux du document pour trouver le placeholder
    for table_obj in doc.tables:
        for row_idx, row in enumerate(table_obj.rows):
            for cell_idx, cell in enumerate(row.cells):
                cell_content = cell.text.strip()
                if has_placeholder(cell_content, placeholder_key):
                    for p in cell.paragraphs:
                        if has_placeholder(p.text, placeholder_key):
                            clear_paragraph_text(p)
                    # Conserve une ligne modèle si elle existe pour préserver le style
                    template_row = table_obj.rows[1] if len(table_obj.rows) > 1 else None
                    while len(table_obj.rows) > 2:
                        table_obj._tbl.remove(table_obj.rows[2]._tr)
                    # Conserve l'en-tête existant du template
                    # Ajoute les lignes du DataFrame
                    for i, (_, row_df) in enumerate(results_sorted.iterrows()):
                        if i == 0 and template_row is not None:
                            row_cells = template_row.cells
                        else:
                            if template_row is not None:
                                new_tr = copy.deepcopy(template_row._tr)
                                table_obj._tbl.append(new_tr)
                                row_cells = table_obj.rows[-1].cells
                            else:
                                row_cells = table_obj.add_row().cells
                        set_cell_text(row_cells[0], row_df['site_name'])
                        set_cell_text(row_cells[1], f"{row_df['score']:.2f}")
                        set_cell_text(row_cells[2], f"{row_df['power_kW']:.2f}")
                        set_cell_text(row_cells[3], f"{row_df['estimated_flow']:.0f}")
                        set_cell_text(row_cells[4], f"{row_df['delta_p']:.2f}")
                    return


def normalize_table_headers(doc: Document) -> None:
    """Ajuste les en-tetes des tableaux sans changer le style du template."""
    for table in doc.tables:
        # Tableau general
        if table.rows and any(has_placeholder(c.text, 'tableau_general') for c in table.rows[1].cells):
            headers = ['Site', 'Score', 'Puissance (kW)', 'Debit (m3/h)', 'Pression (bar)']
            for i, h in enumerate(headers):
                if i < len(table.rows[0].cells):
                    set_cell_text(table.rows[0].cells[i], h)
            # Nettoie les colonnes restantes
            for i in range(len(headers), len(table.rows[0].cells)):
                set_cell_text(table.rows[0].cells[i], '')
        # Synthese sites
        if table.rows and any(has_placeholder(c.text, 'synthese_sites_table') for c in table.rows[1].cells):
            headers = ['Site', 'Score', 'Puissance (kW)', 'Debit (m3/h)', 'Pression (bar)']
            for i, h in enumerate(headers):
                if i < len(table.rows[0].cells):
                    set_cell_text(table.rows[0].cells[i], h)
            for i in range(len(headers), len(table.rows[0].cells)):
                set_cell_text(table.rows[0].cells[i], '')
        # Turbines top 3
        if table.rows and any(has_placeholder(c.text, 'tableau_turbines_top3') for c in table.rows[1].cells):
            headers = ['Type', 'D (mm)', 'Pmin (kW)', 'Pmax (kW)']
            for i, h in enumerate(headers):
                if i < len(table.rows[0].cells):
                    set_cell_text(table.rows[0].cells[i], h)
            for i in range(len(headers), len(table.rows[0].cells)):
                set_cell_text(table.rows[0].cells[i], '')


# -------------------- DATA PROCESSING -------------------- #

def load_and_compute(csv_path: str, turbine_db_path: str = None) -> Tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    raw_dict = load_data({'prv': csv_path})
    raw_df = raw_dict['prv']
    flow_df = compute_hydraulics(raw_df)
    
    # Charger la base de turbines et sélectionner pour avoir le rendement
    turbine_db = load_turbine_db(turbine_db_path)
    flow_df_with_turbines = select_turbine(flow_df, turbine_db, top_n=1)
    
    # Calculer la puissance avec le rendement des turbines sélectionnées
    power_df = compute_power(flow_df_with_turbines)
    score_df = score_sites(flow_df)
    return raw_df, flow_df, power_df, score_df


def prepare_results(flow_df: pd.DataFrame, power_df: pd.DataFrame, score_df: pd.DataFrame) -> pd.DataFrame:
    results = flow_df.copy()
    results['power_kW'] = power_df['power'] / 1000
    results['score'] = score_df['score']
    return results.sort_values('score', ascending=False)


def compute_indicators(results_clean: pd.DataFrame) -> Dict[str, float]:
    return {
        'nb_sites': len(results_clean),
        'delta_p_moy': results_clean['delta_p'].mean(),
        'delta_p_max': results_clean['delta_p'].max(),
        'debit_moy': results_clean['estimated_flow'].mean(),
        'puiss_moy': results_clean['power_kW'].mean(),
        'score_moy': results_clean['score'].mean(),
    }


def build_synthese_top3(results_clean: pd.DataFrame) -> str:
    best_sites = results_clean.head(3)
    synthese = ""
    for _, row in best_sites.iterrows():
        synthese += (
            f"Site: {row['site_name']} | Score: {row['score']:.2f} | "
            f"Puissance: {row['power_kW']:.2f} kW | Débit: {row['estimated_flow']:.0f} m3/h | "
            f"Pression: {row['delta_p']:.2f} bar\n"
        )
    return synthese


def build_quality_summary(raw_df: pd.DataFrame) -> str:
    missing_summary: List[str] = []
    for col in ['estimated_flow', 'delta_p', 'diameter', 'site_name']:
        if col in raw_df.columns:
            missing_summary.append(f"- {col}: {int(raw_df[col].isna().sum())} valeurs manquantes")
    if missing_summary:
        return "Qualite des donnees (comptage simple des manquants)\n" + "\n".join(missing_summary)
    return "Qualite des donnees : aucune valeur manquante detectee sur les champs cles."


# -------------------- BUSINESS LOGIC -------------------- #

def build_section1_content(indicators: Dict[str, float], qualite_donnees: str, synthese_top3: str) -> str:
    return "\n".join([
        "Contexte et objectifs du rapport :",
        "- Valoriser l'energie dissipee dans les reducteurs de pression.",
        "- Identifier les sites a fort potentiel selon des criteres techniques.",
        "- Proposer une selection initiale de turbines compatibles.",
        "\n" + qualite_donnees,
        "\nIndicateurs globaux :",
        f"- Nombre de sites : {indicators['nb_sites']}",
        f"- Delta P moyen : {indicators['delta_p_moy']:.2f} bar (max {indicators['delta_p_max']:.2f} bar)",
        f"- Debit moyen : {indicators['debit_moy']:.1f} m3/h",
        f"- Puissance moyenne : {indicators['puiss_moy']:.2f} kW",
        f"- Score moyen : {indicators['score_moy']:.2f}",
        "\nSynthese des 3 meilleurs sites :",
        synthese_top3.strip() if synthese_top3.strip() else "Aucun site disponible."
    ]).strip()


def build_section2_content(results_clean: pd.DataFrame, turbine_db: pd.DataFrame, top_n: int) -> str:
    hypotheses = (
        "Hypotheses productible/economie\n"
        "- Heures de fonctionnement/an: 6500\n"
        "- Disponibilite turbine: 0.93\n"
        "- Prix de l'electricite: 0.18 EUR/kWh\n"
        "- Consommation site: 10000 kWh/an\n"
        "- Mode: autoconsommation\n"
        "- CAPEX/OPEX: 0\n"
        "- Subvention: 0 % CAPEX\n"
        "- Taux d'actualisation: 5 %\n"
        "- Duree de vie projet: 20 ans"
    )
    sites_details: List[str] = []
    for _, row in results_clean.head(top_n).iterrows():
        candidates = propose_turbines(row, turbine_db, top_n=3)
        if not candidates.empty:
            types = ", ".join([str(t) for t in candidates['type_turbine'].head(3)])
        else:
            types = "Aucune turbine compatible"
        sites_details.append(
            f"Site: {row['site_name']} | Puissance: {row['power_kW']:.2f} kW | "
            f"Debit: {row['estimated_flow']:.0f} m3/h | Pression: {row['delta_p']:.2f} bar | "
            f"Score: {row['score']:.2f} | Turbines: {types}"
        )
    if not sites_details:
        return hypotheses + "\n\nAucun site n'a ete detecte pour le dimensionnement."
    return hypotheses + "\n\n" + "\n".join(sites_details)


def build_focus_site_report(results_sorted: pd.DataFrame, turbine_db: pd.DataFrame) -> str:
    if results_sorted.empty:
        return "Aucun site disponible pour la synthese focus."

    focus_row = results_sorted.iloc[0]
    rank = 1
    total_sites = len(results_sorted)

    candidates = propose_turbines(focus_row, turbine_db, top_n=3)
    selected = candidates.iloc[0].to_dict() if not candidates.empty else {}

    site_input = {
        "delta_p": float(focus_row.get("delta_p", 0)),
        "estimated_flow": float(focus_row.get("estimated_flow", 0)),
        "consommation_kwh": 10000,
        "mode": "autoconsommation",
    }
    if selected:
        selected = dict(selected)
        selected["heures_fonctionnement"] = 6500
        selected["availability"] = 0.93
    finance = {
        "electricity_price": 0.18,
        "injection_tariff": 0.18,
        "capex": 0,
        "opex": 0,
    }
    econ = productible.run_full_model(site_input, selected, finance) if selected else {}

    lines = [
        "1. Contexte, lieu et objectif",
        f"Site focus: {focus_row.get('site_name', 'Site')}.",
        "Objectif: qualifier le potentiel hydro et prioriser une turbine compatible.",
        "Carte OpenStreetMap: a inserer dans le template.",
        "",
        "2. Estimation du productible et classement",
        f"Classement: {rank}/{total_sites}.",
        f"Delta P: {focus_row.get('delta_p', 0):.2f} bar.",
        f"Debit estime: {focus_row.get('estimated_flow', 0):.0f} m3/h.",
        f"Puissance estimee: {focus_row.get('power_kW', 0):.2f} kW.",
        f"Score: {focus_row.get('score', 0):.2f}.",
        "",
        "3. Turbines recommandees et explications",
    ]
    if candidates.empty:
        lines.append("Aucune turbine compatible identifiee pour ce site.")
    else:
        for _, t in candidates.iterrows():
            desc = str(t.get("description", "")).strip()
            source = str(t.get("source", "")).strip()
            line = f"- {t.get('type_turbine', 'Turbine')} (D {t.get('diametre_mm', 0)} mm, {t.get('puissance_min_kw', 0)}-{t.get('puissance_max_kw', 0)} kW)"
            if desc:
                line += f": {desc}"
            if source:
                line += f" (src: {source})"
            lines.append(line)
    tri_val = econ.get('tri')
    tri_text = f"{tri_val * 100:.1f} %" if tri_val is not None else "-"
    lines += [
        "",
        "4. Analyse economique",
        f"Energie annuelle: {econ.get('energie_kwh', 0):.0f} kWh/an.",
        f"Economies: {econ.get('economies_eur', 0):.0f} EUR/an.",
        f"Revenus injection: {econ.get('revenus_injection_eur', 0):.0f} EUR/an.",
        f"Taux autoconsommation: {econ.get('taux_autoconsommation', 0) * 100:.1f} %.",
        f"VAN: {econ.get('van_eur', 0):.0f} EUR.",
        f"TRI: {tri_text}.",
        "",
        "5. Conclusion",
        "Le site est prioritaire si les contraintes hydrauliques et d'exploitation sont confirmees.",
    ]
    return "\n".join(lines).strip()


def build_variables(
    indicators: Dict[str, float],
    synthese_top3: str,
    contenu_titre_1: str,
    contenu_titre_2: str,
    focus_site_report: str,
) -> Dict[str, str]:
    intro = (
        "Ce rapport présente une analyse du potentiel hydroélectrique sur le réseau d'eau potable étudié. "
        "L'objectif est de valoriser l'énergie dissipée dans les réducteurs de pression, "
        "d'identifier les sites à fort potentiel, et de recommander des turbines adaptées. "
        "L'analyse comparative permet de sélectionner les meilleurs sites selon des critères techniques et économiques."
    )
    return {
        "nom_projet": "Larzacqua",
        "client": "Projet de Semestre",
        "date": "30/03/2026",
        "phase_cadrage": "Phase de cadrage du projet...",
        "titre_partie_1": "Partie 1 - Analyse comparative",
        "titre_partie_2": "Sous-partie : Synthese et qualite des donnees",
        "titre_partie_3": "Sous-sections : Indicateurs, Tableaux, Synthese",
        "intro": intro,
        "titre_1": "Analyse comparative",
        "contenu_titre_1": contenu_titre_1,
        "nb_sites": str(int(indicators['nb_sites'])),
        "delta_p_moy": f"{indicators['delta_p_moy']:.2f}",
        "delta_p_max": f"{indicators['delta_p_max']:.2f}",
        "debit_moy": f"{indicators['debit_moy']:.1f}",
        "puiss_moy": f"{indicators['puiss_moy']:.2f}",
        "score_moy": f"{indicators['score_moy']:.2f}",
        "tableau_general": "",
        "synthese_top3": synthese_top3,
        "titre_2": "Estimation par site (dimensionnement)",
        "contenu_titre_2": contenu_titre_2,
        "synthese_sites_table": "",
        "tableau_turbines_top3": "",
        "titre_3": "Recommandations et pedagogie",
        "contenu_titre_3": (
            "Recommandations pedagogiques :\n"
            "- Prioriser les sites avec meilleur score et forte pression.\n"
            "- Verifier les contraintes d'exploitation et de maintenance.\n"
            "- Confirmer la compatibilite hydraulique avant dimensionnement final."
        ),
        "focus_site_report": focus_site_report,
        "conclusion": "Conclusion du rapport...",
        "sources": "Sources et références...",
        "auteur": "Guillaume Siossian",
        "reference_document": "EPF-P4A-LARZACQUA-2026-001",
    }


# -------------------- WORD GENERATION -------------------- #

def fill_document(doc: Document, variables: Dict[str, str]) -> None:
    for p in doc.paragraphs:
        replace_placeholders_in_paragraph(p, variables)
        strip_toc_errors(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, val in variables.items():
                    if key not in ['tableau_general', 'synthese_sites_table', 'tableau_turbines_top3']:
                        for p in cell.paragraphs:
                            replace_placeholders_in_paragraph(p, {key: val})
                            strip_toc_errors(p)

    for section in doc.sections:
        header = section.header
        for p in header.paragraphs:
            replace_placeholders_in_paragraph(p, variables)
            strip_toc_errors(p)
        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_placeholders_in_paragraph(p, variables)
                        strip_toc_errors(p)
        footer = section.footer
        for p in footer.paragraphs:
            replace_placeholders_in_paragraph(p, variables)
            strip_toc_errors(p)
        for table in footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_placeholders_in_paragraph(p, variables)
                        strip_toc_errors(p)


def insert_tables(doc: Document, results_clean: pd.DataFrame, best_sites: pd.DataFrame, turbine_db: pd.DataFrame) -> None:
    normalize_table_headers(doc)
    cols_general = ['site_name', 'score', 'power_kW', 'estimated_flow', 'delta_p']
    tableau_general_df = results_clean[cols_general].copy()
    tableau_general_df['score'] = tableau_general_df['score'].round(2)
    tableau_general_df['power_kW'] = tableau_general_df['power_kW'].round(2)
    tableau_general_df['estimated_flow'] = tableau_general_df['estimated_flow'].round(0).astype(int)
    tableau_general_df['delta_p'] = tableau_general_df['delta_p'].round(2)
    insert_dataframe_to_word(doc, tableau_general_df, 'tableau_general')
    synthese_df = results_clean.copy()
    synthese_df['score'] = synthese_df['score'].round(2)
    synthese_df['power_kW'] = synthese_df['power_kW'].round(2)
    synthese_df['estimated_flow'] = synthese_df['estimated_flow'].round(0).astype(int)
    synthese_df['delta_p'] = synthese_df['delta_p'].round(2)
    insert_synthese_sites_table(doc, synthese_df, 'synthese_sites_table')

    turbines_top3_df = pd.DataFrame()
    for _, row in best_sites.iterrows():
        candidates = turbine_db[(turbine_db['pression_min_bar'] <= row['delta_p']) &
                                (turbine_db['pression_max_bar'] >= row['delta_p']) &
                                (turbine_db['debit_min_m3h'] <= row['estimated_flow']) &
                                (turbine_db['debit_max_m3h'] >= row['estimated_flow'])]
        if not candidates.empty:
            turbines_top3_df = pd.concat([
                turbines_top3_df,
                candidates[['type_turbine', 'diametre_mm', 'puissance_min_kw', 'puissance_max_kw']].head(3)
            ], ignore_index=True)
    if not turbines_top3_df.empty:
        turbines_top3_df = turbines_top3_df.copy()
        turbines_top3_df['type_turbine'] = turbines_top3_df['type_turbine'].astype(str).str.slice(0, 12)
        turbines_top3_df['diametre_mm'] = turbines_top3_df['diametre_mm'].round(0).astype(int)
        turbines_top3_df['puissance_min_kw'] = turbines_top3_df['puissance_min_kw'].round(1)
        turbines_top3_df['puissance_max_kw'] = turbines_top3_df['puissance_max_kw'].round(1)
        insert_dataframe_to_word(doc, turbines_top3_df, 'tableau_turbines_top3')
        for table in doc.tables:
            if table.rows and any(has_placeholder(c.text, 'tableau_turbines_top3') for c in table.rows[1].cells):
                table.alignment = WD_TABLE_ALIGNMENT.LEFT
                table.autofit = False
                # 4 colonnes: Type, D, Pmin, Pmax
                set_table_column_widths(table, [2.2, 1.0, 1.4, 1.4])


# -------------------- PDF EXPORT -------------------- #

def export_pdf(docx_path: str, pdf_path: str) -> None:
    strip_comments_in_docx(docx_path)
    strip_vml_shapes_in_docx(docx_path)
    convert(docx_path, pdf_path)


# -------------------- MAIN -------------------- #

def main() -> None:
    template_path = PROVIDED_TEMPLATE if docx_has_images(PROVIDED_TEMPLATE) else DEFAULT_TEMPLATE
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    raw_df, flow_df, power_df, score_df = load_and_compute(CSV_PATH, TURBINE_DB_PATH)
    results_sorted = prepare_results(flow_df, power_df, score_df)
    results_clean = results_sorted.dropna(subset=['site_name', 'estimated_flow', 'delta_p', 'power_kW', 'score']).copy()
    indicators = compute_indicators(results_clean)
    synthese_top3 = build_synthese_top3(results_clean)
    qualite_donnees = build_quality_summary(raw_df)
    contenu_titre_1 = build_section1_content(indicators, qualite_donnees, synthese_top3)
    turbine_db = load_turbine_db(TURBINE_DB_PATH)
    contenu_titre_2 = build_section2_content(results_clean, turbine_db, TOP_N_SITES)
    focus_site_report = build_focus_site_report(results_sorted, turbine_db)
    variables = build_variables(indicators, synthese_top3, contenu_titre_1, contenu_titre_2, focus_site_report)

    doc = Document(template_path)
    for section in doc.sections:
        section.footer_distance = Inches(0.1)

    fill_document(doc, variables)
    insert_tables(doc, results_clean, results_clean.head(3), turbine_db)
    doc.save(TEMP_DOCX)
    export_pdf(TEMP_DOCX, OUTPUT_PDF)
    print(f"PDF généré fidèlement au template Word : {OUTPUT_PDF}")

if __name__ == "__main__":
    main()